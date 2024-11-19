import os
import pandas as pd
import win32com.client as win32
from docx import Document
from fpdf import FPDF

# File paths
CERTIFICATES_PATH = r"C:\Users\User\Downloads\certs"
EXCEL_FILE = r"C:\Users\User\PycharmProjects\Bulk email\delegates.xlsx"

# Email content
EMAIL_SUBJECT = "Certificate of Attendance for GCC Conference"
EMAIL_BODY_TEMPLATE = """\
Dear {name},

Thank you for attending the GCC Conference. We appreciate your participation and hope you found the sessions insightful.

Please find attached your Certificate of Attendance.

Kind Regards,  
GCC Team
"""

# Logs
success_log = []
failure_log = []


# Function to get available Outlook accounts
def get_outlook_accounts():
    outlook = win32.Dispatch("Outlook.Application").Session
    accounts = []
    for account in outlook.Accounts:
        accounts.append(account.DisplayName)
    return accounts


# Function to send an email with attachment using a specific Outlook account
def send_email(account_name, to_email, subject, body, attachment_path):
    try:
        outlook = win32.Dispatch("Outlook.Application")
        mail = outlook.CreateItem(0)
        mail.Subject = subject
        mail.Body = body
        mail.To = to_email
        mail.Attachments.Add(attachment_path)

        # Set the account to send from
        mail.SendUsingAccount = next(
            account for account in outlook.Session.Accounts if account.DisplayName == account_name
        )
        mail.Send()
        return True
    except Exception as e:
        print(f"Failed to send email to {to_email}: {e}")
        return False


# Function to generate Word report
def generate_word_report(success_log, failure_log):
    doc = Document()
    doc.add_heading("Email Sending Report", level=1)

    # Add success section
    doc.add_heading("Successfully Sent Emails", level=2)
    for record in success_log:
        doc.add_paragraph(f"{record['name']} - {record['email']}")

    # Add failure section
    doc.add_heading("Failed Emails", level=2)
    for record in failure_log:
        doc.add_paragraph(f"{record['name']} - {record['email']} - {record['reason']}")

    report_path = os.path.join(os.getcwd(), "Email_Report.docx")
    doc.save(report_path)
    print(f"Word report saved to {report_path}")


# Function to generate PDF report
def generate_pdf_report(success_log, failure_log):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # Add title
    pdf.set_font("Arial", size=16, style="B")
    pdf.cell(0, 10, "Email Sending Report", ln=True, align="C")

    # Add success section
    pdf.set_font("Arial", size=14, style="B")
    pdf.cell(0, 10, "Successfully Sent Emails:", ln=True)
    pdf.set_font("Arial", size=12)
    for record in success_log:
        pdf.cell(0, 10, f"{record['name']} - {record['email']}", ln=True)

    # Add failure section
    pdf.set_font("Arial", size=14, style="B")
    pdf.cell(0, 10, "Failed Emails:", ln=True)
    pdf.set_font("Arial", size=12)
    for record in failure_log:
        pdf.cell(0, 10, f"{record['name']} - {record['email']} - {record['reason']}", ln=True)

    report_path = os.path.join(os.getcwd(), "Email_Report.pdf")
    pdf.output(report_path)
    print(f"PDF report saved to {report_path}")


# Main function to process the Excel file and send emails
def main():
    # Get list of Outlook accounts
    accounts = get_outlook_accounts()
    if not accounts:
        print("No Outlook accounts found. Please configure Outlook.")
        return

    print("Available Outlook Accounts:")
    for idx, account in enumerate(accounts):
        print(f"{idx + 1}. {account}")

    selected_index = int(input("Select the account to use (1, 2, etc.): ")) - 1
    if selected_index < 0 or selected_index >= len(accounts):
        print("Invalid selection.")
        return

    selected_account = accounts[selected_index]
    print(f"Using account: {selected_account}")

    # Read the Excel file
    data = pd.read_excel(EXCEL_FILE)

    # Loop through the rows and send emails
    for _, row in data.iterrows():
        name = row["Name"].strip()
        email = row["Email"].strip()
        certificate_path = os.path.join(CERTIFICATES_PATH, f"{name}.pdf")

        if os.path.exists(certificate_path):
            print(f"Sending certificate to {name} at {email}...")
            email_body = EMAIL_BODY_TEMPLATE.format(name=name)
            success = send_email(
                account_name=selected_account,
                to_email=email,
                subject=EMAIL_SUBJECT,
                body=email_body,
                attachment_path=certificate_path,
            )
            if success:
                success_log.append({"name": name, "email": email})
            else:
                failure_log.append({"name": name, "email": email, "reason": "Email sending failed"})
        else:
            print(f"Certificate file not found: {certificate_path}")
            failure_log.append({"name": name, "email": email, "reason": "Certificate not found"})

    # Generate reports
    generate_word_report(success_log, failure_log)
    generate_pdf_report(success_log, failure_log)


# Run the main function automatically
if __name__ == "__main__":
    main()
